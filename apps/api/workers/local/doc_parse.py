"""In-process document parsing.

Replaces the RebaseKit ``/docparse/api/parse`` call. Supports PDF,
DOCX, XLSX, and plain text. Accepts either a URL (fetched via httpx)
or base64-encoded content.
"""

from __future__ import annotations

import base64
import io
import ssl
from dataclasses import dataclass
from typing import Any

import httpx

_MAX_DOCUMENT_BYTES = 20 * 1024 * 1024  # 20 MB hard cap
_FETCH_TIMEOUT = 20.0

# Use the OS trust store instead of httpx's bundled certifi so we pick
# up CA chains that certifi's Mozilla bundle occasionally misses — see
# workers/local/llm_tasks.py for the full rationale.
_SSL_CONTEXT = ssl.create_default_context()


class DocumentParseError(ValueError):
    """Raised on bad input or unparseable content."""


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    text: str
    page_count: int
    format: str
    metadata: dict[str, Any]


def _format_from_bytes(data: bytes, mime_type: str | None) -> str:
    """Guess the document format from magic bytes first, mime second."""
    if data.startswith(b"%PDF-"):
        return "pdf"
    # DOCX / XLSX are ZIP files with specific content types
    if data[:2] == b"PK":
        if mime_type and "wordprocessing" in mime_type:
            return "docx"
        if mime_type and ("spreadsheet" in mime_type or "excel" in mime_type):
            return "xlsx"
        # Fall through to ZIP inspection
        return _zip_format(data)
    if mime_type:
        if "pdf" in mime_type:
            return "pdf"
        if "wordprocessing" in mime_type or "docx" in mime_type:
            return "docx"
        if "spreadsheet" in mime_type or "excel" in mime_type or "xlsx" in mime_type:
            return "xlsx"
        if mime_type.startswith("text/"):
            return "text"
    return "text"  # best-effort fallback


def _zip_format(data: bytes) -> str:
    """Peek into a ZIP archive to distinguish DOCX vs XLSX."""
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            names = set(z.namelist())
    except zipfile.BadZipFile:
        return "text"
    if "word/document.xml" in names:
        return "docx"
    if any(n.startswith("xl/") for n in names):
        return "xlsx"
    return "text"


def _parse_pdf(data: bytes) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — pypdf raises on malformed pages
            pages.append("")
    text = "\n\n".join(pages).strip()
    meta = {}
    try:
        if reader.metadata:
            meta = {
                k.lstrip("/"): str(v)
                for k, v in reader.metadata.items()
            }
    except Exception:  # noqa: BLE001
        meta = {}
    return ParsedDocument(
        text=text,
        page_count=len(reader.pages),
        format="pdf",
        metadata=meta,
    )


def _parse_docx(data: bytes) -> ParsedDocument:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return ParsedDocument(
        text="\n".join(parts).strip(),
        page_count=0,
        format="docx",
        metadata={},
    )


def _parse_xlsx(data: bytes) -> ParsedDocument:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet_name in wb.sheetnames:
        lines.append(f"# Sheet: {sheet_name}")
        sheet = wb[sheet_name]
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                lines.append(
                    "\t".join("" if c is None else str(c) for c in row)
                )
    wb.close()
    return ParsedDocument(
        text="\n".join(lines).strip(),
        page_count=len(wb.sheetnames),
        format="xlsx",
        metadata={"sheets": wb.sheetnames},
    )


def _parse_text(data: bytes) -> ParsedDocument:
    text = data.decode("utf-8", errors="replace")
    return ParsedDocument(
        text=text,
        page_count=1,
        format="text",
        metadata={},
    )


def parse_bytes(data: bytes, mime_type: str | None = None) -> ParsedDocument:
    """Dispatch to the correct parser based on format detection."""
    if len(data) > _MAX_DOCUMENT_BYTES:
        raise DocumentParseError(
            f"document exceeds the {_MAX_DOCUMENT_BYTES // (1024 * 1024)} MB limit"
        )
    fmt = _format_from_bytes(data, mime_type)
    if fmt == "pdf":
        return _parse_pdf(data)
    if fmt == "docx":
        return _parse_docx(data)
    if fmt == "xlsx":
        return _parse_xlsx(data)
    return _parse_text(data)


async def _fetch(url: str) -> tuple[bytes, str | None]:
    async with httpx.AsyncClient(
        timeout=_FETCH_TIMEOUT,
        follow_redirects=True,
        verify=_SSL_CONTEXT,
    ) as client:
        r = await client.get(url)
        r.raise_for_status()
        content_type = r.headers.get("content-type", "").split(";")[0].strip() or None
        return r.content, content_type


async def run(inp: dict) -> dict:
    """Task-handler entry point."""
    url = inp.get("url")
    content_b64 = inp.get("content_base64") or inp.get("base64_content")

    if url:
        try:
            data, mime_type = await _fetch(url)
        except (httpx.HTTPError, OSError) as exc:
            raise DocumentParseError(f"failed to fetch document: {exc}") from exc
    elif content_b64:
        try:
            data = base64.b64decode(content_b64, validate=False)
        except (ValueError, TypeError) as exc:
            raise DocumentParseError(f"invalid base64 content: {exc}") from exc
        mime_type = inp.get("mime_type")
    else:
        raise DocumentParseError(
            "document_parse requires either 'url' or 'content_base64'"
        )

    parsed = parse_bytes(data, mime_type=mime_type)
    return {
        "text": parsed.text,
        "format": parsed.format,
        "page_count": parsed.page_count,
        "metadata": parsed.metadata,
    }


__all__ = [
    "DocumentParseError",
    "ParsedDocument",
    "parse_bytes",
    "run",
]
